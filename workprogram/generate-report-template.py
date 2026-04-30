#!/usr/bin/env python3
"""
Generates two NACSA NCII audit report documents:
1. NACSA-Audit-Report-Template.docx — blank report template for auditors
2. NACSA-Sample-Audit-Report.docx — fictional example for reference

Aligned to:
- CE Directive 08 para 23 submission requirements
- Control Library v4.0 (126 controls, 210 Tier 1, 18 Tier 2, 42 directive requirements)
- CoP Template v1.6 (18 domains, sections 4.0-21.0)
- Two-tier requirement structure (Tier 1 = findings, Tier 2 = improvement opportunities)

Usage: python3 generate-report-template.py
"""

import os
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LIBRARY_PATH = os.path.join(SCRIPT_DIR, "..", "controls", "library.json")
DOMAINS_PATH = os.path.join(SCRIPT_DIR, "..", "controls", "domains.json")

# --- Colour palette ---
NAVY = "1F4E79"
DARK_RED = RGBColor(0x8B, 0x00, 0x00)
NAVY_RGB = RGBColor(0x1F, 0x4E, 0x79)
LIGHT_BLUE = "D6E4F0"
LIGHT_GREY = "F2F2F2"
LIGHT_RED = "FFC7CE"
LIGHT_YELLOW = "FFEB9C"
LIGHT_GREEN = "C6EFCE"
WHITE = "FFFFFF"

# --- Domain data (v4.0 aligned to CoP Template v1.6) ---
DOMAIN_DATA = [
    ("4.0", "Governance and Leadership", "governance-leadership", 1, 1, 0),
    ("5.0", "Cyber Security Policy and Objectives", "cyber-security-policy", 8, 11, 4),
    ("6.0", "Organisational Development", "organisational-development", 10, 19, 2),
    ("7.0", "Cyber Security Assurance", "cyber-security-assurance", 11, 24, 2),
    ("8.0", "Resource Allocation and Optimisation", "resource-allocation", 8, 19, 0),
    ("9.0", "Risk Management", "risk-management", 9, 19, 2),
    ("10.0", "Operational Efficiency", "operational-efficiency", 4, 9, 0),
    ("11.0", "Data Security", "data-security", 15, 23, 0),
    ("12.0", "Contractual Management", "contractual-management", 3, 6, 0),
    ("13.0", "Physical Security", "physical-security", 8, 15, 1),
    ("14.0", "System and Network Security", "system-network-security", 16, 24, 0),
    ("15.0", "Access Control", "access-control", 11, 13, 3),
    ("16.0", "Technical Vulnerability", "technical-vulnerability", 9, 10, 1),
    ("17.0", "Cyber Security Event Management", "event-management", 3, 3, 0),
    ("18.0", "Cyber Security Incident Management", "incident-management", 3, 6, 0),
    ("19.0", "Business Continuity Management", "business-continuity", 5, 6, 3),
    ("20.0", "Sector-Specific Requirements", "sector-specific", 2, 2, 0),
    ("21.0", "Monitoring and Compliance", "monitoring-compliance", 0, 0, 0),
]

# 11 statutory obligations (Act 854 Part IV)
STATUTORY_OBLIGATIONS = [
    ("s17", "Designation as NCII Entity", "Perlantikan sebagai Entiti IKMN"),
    ("s18", "General Duties of NCII Entity", "Kewajipan Am Entiti IKMN"),
    ("s19", "Notification of Change of Ownership/Control", "Pemberitahuan Perubahan Pemilikan/Kawalan"),
    ("s20", "Provision of Information", "Pemberian Maklumat"),
    ("s21", "Risk Assessment", "Penilaian Risiko"),
    ("s22", "Audit and Code of Practice Compliance", "Audit dan Pematuhan Tataamalan"),
    ("s23", "Audit Report Submission (30 days)", "Penghantaran Laporan Audit (30 hari)"),
    ("s24", "Compliance with Standards", "Pematuhan Piawaian"),
    ("s25", "Cyber Threat Information Sharing", "Perkongsian Maklumat Ancaman Siber"),
    ("s26", "Incident Notification (6-hour)", "Pemberitahuan Insiden (6 jam)"),
    ("s27", "Compliance with CE Directives", "Pematuhan Arahan Ketua Eksekutif"),
]

# --- Helper functions ---

def set_cell_shading(cell, color):
    """Set cell background colour."""
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")


def add_styled_table(doc, headers, rows, col_widths=None, header_color=NAVY,
                     alt_row=True, font_size=9):
    """Add a formatted table with navy headers and optional alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
        if alt_row and r % 2 == 1:
            for c in range(len(headers)):
                set_cell_shading(table.rows[r + 1].cells[c], LIGHT_GREY)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_severity_badge(doc, severity, text=None):
    """Add a severity-coloured paragraph."""
    colours = {
        "Kritikal": LIGHT_RED,
        "Tinggi": LIGHT_YELLOW,
        "Sederhana": LIGHT_YELLOW,
        "Rendah": LIGHT_GREEN,
    }
    p = doc.add_paragraph()
    run = p.add_run(text or severity)
    run.bold = True
    run.font.size = Pt(10)
    return p


def add_finding_table(doc, fields, col_widths=None):
    """Add a finding detail table with the specific formatting."""
    if col_widths is None:
        col_widths = [5, 11]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r, (label, value) in enumerate(fields):
        label_cell = table.rows[r].cells[0]
        value_cell = table.rows[r].cells[1]
        label_cell.text = label
        value_cell.text = str(value)
        # Style label
        for p in label_cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(label_cell, LIGHT_BLUE)
        # Style value
        for p in value_cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

    for row in table.rows:
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].width = Cm(col_widths[1])

    return table


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)


def add_cover_page(doc, entity_data):
    """Add the cover page."""
    for _ in range(3):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAPORAN AUDIT KESELAMATAN SIBER\nCYBER SECURITY AUDIT REPORT")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY_RGB

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Akta Keselamatan Siber 2024 (Akta 854)\nCyber Security Act 2024 (Act 854)")
    run.font.size = Pt(14)

    doc.add_paragraph("")

    cover_fields = [
        ("Nama Entiti NCII / NCII Entity Name", entity_data.get("entity_name", "[Entity Name]")),
        ("Sektor NCII / NCII Sector", entity_data.get("sector", "[Sector]")),
        ("Ketua Sektor / Sector Lead", entity_data.get("sector_lead", "[Sector Lead]")),
        ("Bil. Sistem NCII / NCII Systems Count", entity_data.get("system_count", "[N]")),
        ("Tempoh Penilaian / Assessment Period", entity_data.get("period", "[Start Date] to [End Date]")),
        ("Tarikh Laporan / Report Date", entity_data.get("report_date", "[DD/MM/YYYY]")),
        ("Firma Audit / Audit Firm", entity_data.get("audit_firm", "[Firm Name]")),
        ("Rakan Kongsi Jaminan / Engagement Partner", entity_data.get("partner", "[Name, Qualification]")),
        ("Ketua Juruaudit / Lead Auditor", entity_data.get("lead_auditor", "[Name, Qualification]")),
        ("Pengamal Undang-Undang / Legal Practitioner (Mode A)", entity_data.get("legal", "[Name, Firm] or N/A")),
        ("No. Rujukan / Reference Number", entity_data.get("ref_no", "[NACSA-AUDIT-YYYY-NNN]")),
    ]
    add_styled_table(doc, ["Butiran / Field", "Nilai / Value"], cover_fields, col_widths=[8, 10])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SULIT / CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_RED


def add_distribution(doc, is_sample=False):
    """Distribution and document control section."""
    doc.add_page_break()
    doc.add_heading("Pengedaran dan Kawalan Dokumen / Distribution and Document Control", level=1)

    doc.add_heading("Senarai Pengedaran / Distribution List", level=2)
    if is_sample:
        dist = [
            ("1", "Ketua Eksekutif NACSA / Chief Executive", "NACSA", "1"),
            ("2", "Suruhanjaya Tenaga / Energy Commission", "Ketua Sektor", "2"),
            ("3", "Pengurusan Syarikat Tenaga Siber", "Entiti NCII", "3"),
            ("4", "Jawatankuasa Audit", "Syarikat Tenaga Siber", "4"),
        ]
    else:
        dist = [
            ("1", "Ketua Eksekutif NACSA / Chief Executive", "NACSA", "1"),
            ("2", "Ketua Sektor / Sector Lead", "[Sector Lead]", "2"),
            ("3", "Pengurusan Entiti / Entity Management", "[Entity]", "3"),
            ("4", "Jawatankuasa Audit / Audit Committee", "[Entity]", "4"),
        ]
    add_styled_table(doc,
        ["#", "Penerima / Recipient", "Organisasi / Organisation", "Salinan / Copy #"],
        dist, col_widths=[1, 6, 5, 2])

    doc.add_paragraph("")
    doc.add_heading("Sejarah Versi / Version History", level=2)
    if is_sample:
        versions = [
            ("0.1", "10/02/2026", "Ahmad Faizal", "Draf untuk semakan dalaman"),
            ("0.2", "17/02/2026", "Ahmad Faizal", "Semakan pengamal undang-undang"),
            ("1.0", "24/02/2026", "Ahmad Faizal", "Muktamad — dihantar kepada KE NACSA"),
        ]
    else:
        versions = [
            ("0.1", "[Date]", "[Author]", "Draft for internal review"),
            ("0.2", "[Date]", "[Author]", "Revised after legal practitioner review"),
            ("1.0", "[Date]", "[Author]", "Final — submitted to NACSA CE"),
        ]
    add_styled_table(doc,
        ["Versi / Version", "Tarikh / Date", "Penulis / Author", "Perihalan / Description"],
        versions, col_widths=[2, 3, 4, 7])


def add_toc(doc, include_directive_section=True):
    """Table of contents."""
    doc.add_page_break()
    doc.add_heading("Jadual Kandungan / Table of Contents", level=1)
    toc_items = [
        "1. Ringkasan Eksekutif / Executive Summary",
        "2. Skop dan Objektif Audit / Audit Scope and Objectives",
        "3. Pendekatan Audit / Audit Approach and Methodology",
        "4. Pematuhan Obligasi Berkanun / Statutory Obligation Compliance",
        "5. Penilaian Liabiliti Pengarah (s52) / Director Liability Assessment",
        "6. Penemuan Mengikut Domain CoP / Findings by CoP Domain (Tier 1)",
        "7. Pematuhan Arahan Ketua Eksekutif / CE Directive Compliance",
        "8. Profil Pematuhan Domain / Domain Compliance Profile",
        "9. Ringkasan Penemuan / Finding Summary",
        "10. Peluang Penambahbaikan / Improvement Opportunities (Tier 2)",
        "11. Rumusan / Conclusion",
        "12. Pelan Tindakan Pengurusan / Management Action Plan",
        "Lampiran A — Senarai Sistem NCII / NCII System Register",
        "Lampiran B — Pendedahan Penalti / Penalty Exposure Summary",
        "Lampiran C — Peringkat Keterukan / Severity Definitions",
        "Lampiran D — Glosari / Glossary",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")


def add_executive_summary(doc, data=None):
    """Section 1: Executive Summary."""
    doc.add_page_break()
    doc.add_heading("1. Ringkasan Eksekutif / Executive Summary", level=1)

    if data:
        doc.add_paragraph(data.get("intro", ""))
    else:
        p = doc.add_paragraph(
            "This section provides a high-level summary of the audit findings, overall compliance posture, "
            "and key areas requiring immediate attention. It should be readable by board members and senior "
            "management without reference to the detailed findings."
        )
        p.italic = True

    doc.add_heading("1.1 Keseluruhan Postur Pematuhan / Overall Compliance Posture", level=2)
    if data:
        metrics = data.get("metrics", [])
    else:
        metrics = [
            ("Total Tier 1 CoP requirements assessed", "[N] of 210"),
            ("Directive requirements assessed", "[N] of 42"),
            ("Patuh (Compliant)", "[N] ([%])"),
            ("Separa Patuh (Partially Compliant)", "[N] ([%])"),
            ("Tidak Patuh (Non-Compliant)", "[N] ([%])"),
            ("Tidak Berkenaan (Not Applicable)", "[N]"),
            ("Tier 2 improvement opportunities identified", "[N] of 18"),
            ("Pendapat Tambahan / Supplementary Opinion",
             "[Memuaskan / Perlu Penambahbaikan / Tidak Memuaskan]"),
        ]
    add_styled_table(doc, ["Metrik / Metric", "Nilai / Value"], metrics, col_widths=[8, 8])

    doc.add_paragraph("")
    doc.add_heading("1.2 Ringkasan Penemuan / Finding Summary", level=2)
    if data:
        findings_summary = data.get("findings_summary", [])
    else:
        findings_summary = [
            ("Kritikal (Critical)", "[N]", "[Brief description]"),
            ("Tinggi (High)", "[N]", "[Brief description]"),
            ("Sederhana (Medium)", "[N]", "[Brief description]"),
            ("Rendah (Low)", "[N]", "[Brief description]"),
            ("Jumlah / Total", "[N]", ""),
        ]
    add_styled_table(doc,
        ["Keterukan / Severity", "Bilangan / Count", "Bidang Utama / Key Areas"],
        findings_summary, col_widths=[4, 2, 10])

    doc.add_paragraph("")
    doc.add_heading("1.3 Perkara Memerlukan Perhatian Segera / Key Areas Requiring Immediate Attention", level=2)
    if data:
        for item in data.get("immediate_attention", []):
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph(
            "CE Directive 08 para 23(j): \"Rumusan perlu menekankan aspek utama yang memerlukan "
            "perhatian dan tindakan segera.\""
        ).italic = True
        doc.add_paragraph("[List Critical and High findings requiring immediate remediation]")
        doc.add_paragraph("[Identify systemic patterns across multiple domains]")


def add_scope_section(doc, data=None):
    """Section 2: Scope and Objectives."""
    doc.add_page_break()
    doc.add_heading("2. Skop dan Objektif Audit / Audit Scope and Objectives", level=1)

    doc.add_heading("2.1 Objektif Audit / Audit Objectives", level=2)
    doc.add_paragraph(
        "To determine the extent to which the NCII entity complies with the audit criteria "
        "(\"sejauh mana kriteria audit dipenuhi\") per CE Directive 08 para 5(a)."
    )

    doc.add_heading("2.2 Kriteria Audit / Audit Criteria", level=2)
    criteria = [
        "Cyber Security Act 2024 (Act 854) — Part IV (s17-s27)",
        "Regulations under Act 854: P.U.(A) 219-222/2024",
        "CE Directives (Arahan Ketua Eksekutif) 01-10",
        "Code of Practice (Tataamalan) — CoP Template v1.6, sections 4.0-21.0",
        "[Sector-specific CoP if applicable and verified with sector lead]",
    ]
    for c in criteria:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("2.3 Skop / Scope", level=2)
    if data:
        scope_rows = data.get("scope_rows", [])
    else:
        scope_rows = [
            ("NCII Systems in Scope", "[N systems — see Lampiran A]"),
            ("Tier 1 CoP Requirements", "210 (SHALL clauses from 126 controls across 17 active domains)"),
            ("Tier 2 CoP Requirements", "18 (SHOULD/MAY clauses — improvement opportunities only)"),
            ("CE Directive Requirements", "42 (from CE Directives 01-10)"),
            ("Total Testable Requirements", "210 Tier 1 + 42 directive = 252 compliance-rated"),
            ("CoP Domains Assessed", "4.0-21.0 (18 domains per CoP Template v1.6)"),
            ("Assessment Period", "[Start Date] to [End Date]"),
            ("Exclusions", "[Any domains or systems excluded with justification]"),
            ("Scope Limitations", "[Any limitations encountered]"),
        ]
    add_styled_table(doc,
        ["Elemen Skop / Scope Element", "Perihalan / Description"],
        scope_rows, col_widths=[5, 11])

    doc.add_paragraph("")
    doc.add_heading("2.4 Pasukan Audit / Audit Team", level=2)
    doc.add_paragraph(
        "Per CE Directive 08 para 16(c)(i): every audit team must include at least one member "
        "with technical expertise in the sector being audited."
    ).italic = True

    if data:
        team = data.get("team", [])
    else:
        team = [
            ("Engagement Partner", "[Name]", "[CISA/CISSP/etc.]", "Overall direction, QA, NACSA submission"),
            ("Lead Auditor (Ketua Juruaudit)", "[Name]", "[CISA/CISSP/etc.]", "Full CoP scope (4.0-21.0)"),
            ("Auditor (Juruaudit)", "[Name]", "[Qualification]", "[Assigned domains]"),
            ("Legal Practitioner (Mode A)", "[Name]", "[Malaysian Bar]", "Statutory obligations, s52, penalty exposure"),
            ("Sector Specialist", "[Name] or N/A", "[Qualification]", "CoP 20.0 sector-specific requirements"),
        ]
    add_styled_table(doc,
        ["Peranan / Role", "Nama / Name", "Kelayakan / Qualification", "Skop / Scope"],
        team, col_widths=[4, 3, 3, 6])


def add_methodology_section(doc, data=None):
    """Section 3: Methodology."""
    doc.add_page_break()
    doc.add_heading("3. Pendekatan Audit / Audit Approach and Methodology", level=1)

    doc.add_heading("3.1 Pendekatan Wajib / Mandatory Approaches", level=2)
    doc.add_paragraph(
        "Compliance-Based (Pendekatan Berasaskan Pematuhan) — assesses extent of compliance with "
        "Act 854, regulations, CE directives, and CoP. Produces compliance status per requirement: "
        "Patuh / Separa Patuh / Tidak Patuh / Tidak Berkenaan."
    )
    doc.add_paragraph(
        "Risk-Based (Pendekatan Berasaskan Risiko) — assesses threat, vulnerability, and impact for "
        "each non-compliance. Produces severity rating per finding: Kritikal / Tinggi / Sederhana / Rendah."
    )

    doc.add_heading("3.2 Pendekatan Pilihan / Optional Approaches Elected", level=2)
    if data:
        approaches = data.get("optional_approaches", [])
    else:
        approaches = [
            ("Control-Based (Kawalan)", "[Elected / Not Elected]", "[Applicable domains]"),
            ("Technical Testing (Ujian Teknikal)", "[Elected / Not Elected]", "[Applicable domains]"),
            ("Inspection & Verification (Pemeriksaan)", "[Elected / Not Elected]", "[Applicable domains]"),
            ("Continuous Improvement (Penambahbaikan)", "[Elected / Not Elected]", "[Applicable domains]"),
        ]
    add_styled_table(doc,
        ["Pendekatan / Approach", "Status", "Domain"],
        approaches, col_widths=[5, 4, 7])

    doc.add_paragraph("")
    doc.add_heading("3.3 Struktur Keperluan Dua Peringkat / Two-Tier Requirement Structure", level=2)
    tier_desc = [
        ("Tier 1 (SHALL)", "210 requirements", "Mandatory compliance. Non-compliance = finding with severity rating. "
         "Compliance rated: Patuh / Separa Patuh / Tidak Patuh / Tidak Berkenaan."),
        ("Tier 2 (SHOULD/MAY)", "18 requirements", "Advisory. Non-compliance = improvement opportunity (peluang penambahbaikan). "
         "No severity rating. No remediation obligation."),
        ("Directive Requirements", "42 requirements", "From CE Directives 01-10. Mandatory compliance. "
         "Tested in separate directive compliance section."),
    ]
    add_styled_table(doc,
        ["Peringkat / Tier", "Bilangan / Count", "Perihalan / Description"],
        tier_desc, col_widths=[3, 3, 10])

    doc.add_paragraph("")
    doc.add_heading("3.4 Definisi Penarafan Pematuhan / Compliance Rating Definitions", level=2)
    ratings = [
        ("P — Patuh (Compliant)", "Requirement fully met: policy exists, approved, applied, current, evidenced."),
        ("SP — Separa Patuh (Partially Compliant)", "Some elements exist but gaps remain; started but not completed."),
        ("TP — Tidak Patuh (Non-Compliant)", "Not met or no evidence; entity not aware of requirement."),
        ("TB — Tidak Berkenaan (Not Applicable)", "Not applicable to entity — written justification required and documented."),
    ]
    add_styled_table(doc, ["Penarafan / Rating", "Definisi / Definition"], ratings, col_widths=[5, 11])

    doc.add_paragraph("")
    doc.add_heading("3.5 Definisi Keterukan Penemuan / Finding Severity Definitions", level=2)
    doc.add_paragraph("See Lampiran C for full definitions. Severity is determined using the risk-based approach "
                       "(threat x vulnerability x impact).")

    doc.add_heading("3.6 Metodologi Persampelan / Sampling Methodology", level=2)
    if data:
        doc.add_paragraph(data.get("sampling", ""))
    else:
        doc.add_paragraph(
            "Judgmental (non-statistical) sampling applied. Sample sizes driven by population size and "
            "control criticality. For system-level controls, each NCII system tested independently."
        )


def add_statutory_section(doc, data=None):
    """Section 4: Statutory Obligations."""
    doc.add_page_break()
    doc.add_heading("4. Pematuhan Obligasi Berkanun / Statutory Obligation Compliance", level=1)
    doc.add_paragraph(
        "Assessment of the entity's compliance with 11 mandatory obligations under Act 854 s17-s27. "
        "This section is jointly assessed by the technical audit team and the legal practitioner (Mode A)."
    ).italic = True

    if data:
        stat_rows = data.get("obligations", [])
    else:
        stat_rows = []
        for sec, eng, malay in STATUTORY_OBLIGATIONS:
            stat_rows.append((sec, f"{eng}\n{malay}", "[Met / Not Met]", "[Notes]"))

    add_styled_table(doc,
        ["Seksyen / Section", "Obligasi / Obligation", "Status", "Nota / Notes"],
        stat_rows, col_widths=[1.5, 6, 2.5, 6])


def add_director_liability_section(doc, data=None):
    """Section 5: Director Liability (s52)."""
    doc.add_page_break()
    doc.add_heading("5. Penilaian Liabiliti Pengarah (s52) / Director Liability Assessment", level=1)
    doc.add_paragraph(
        "Under s52 of Act 854, where a body corporate commits an offence under the Act, "
        "every director, chief executive officer, chief operating officer, manager, secretary, "
        "or other similar officer shall be deemed to have committed that offence unless they can "
        "prove the offence was committed without their consent or connivance, and that they exercised "
        "due diligence to prevent the commission of the offence."
    ).italic = True

    if data:
        doc.add_heading("5.1 Penilaian / Assessment", level=2)
        for item in data.get("assessment", []):
            doc.add_paragraph(item, style="List Bullet")
        doc.add_heading("5.2 Pegawai Berkaitan / Relevant Officers", level=2)
        if data.get("officers"):
            add_styled_table(doc,
                ["Jawatan / Position", "Nama / Name", "Tahap Risiko / Risk Level", "Nota / Notes"],
                data["officers"], col_widths=[4, 4, 3, 5])
    else:
        doc.add_heading("5.1 Penilaian / Assessment", level=2)
        doc.add_paragraph("[Assess whether any non-compliance findings trigger s52 director liability exposure]")
        doc.add_paragraph("[Identify which statutory offences (s17-s27 failures) have been established]")
        doc.add_paragraph("[Assess whether management demonstrated due diligence to prevent the offence]")

        doc.add_heading("5.2 Pegawai Berkaitan / Relevant Officers", level=2)
        add_styled_table(doc,
            ["Jawatan / Position", "Nama / Name", "Tahap Risiko / Risk Level", "Nota / Notes"],
            [
                ("Chief Executive Officer", "[Name]", "[High/Medium/Low]", "[Due diligence evidence]"),
                ("Chief Information Security Officer", "[Name]", "[High/Medium/Low]", "[Due diligence evidence]"),
                ("Chief Technology Officer", "[Name]", "[High/Medium/Low]", "[Due diligence evidence]"),
            ],
            col_widths=[4, 4, 3, 5])

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: This assessment covers Mode A (factual findings included in audit report). "
        "The full Mode B privileged legal opinion on director liability is a separate, "
        "confidential deliverable to the audit committee."
    ).italic = True


def add_findings_section(doc, findings=None):
    """Section 6: Findings by CoP Domain (Tier 1)."""
    doc.add_page_break()
    doc.add_heading("6. Penemuan Mengikut Domain CoP / Findings by CoP Domain (Tier 1)", level=1)
    doc.add_paragraph(
        "This section presents Tier 1 (SHALL) findings organised by CoP domain (4.0-21.0). "
        "Each finding uses the 5 C's framework per CE Directive 08 para 20: "
        "Condition, Criteria, Cause, Consequence, and Corrective Action."
    ).italic = True

    if findings:
        # Group findings by domain
        current_domain = None
        domain_idx = 0
        for f in findings:
            if f["domain"] != current_domain:
                current_domain = f["domain"]
                domain_idx += 1
                doc.add_heading(
                    f"6.{domain_idx} CoP {f['domain_section']} — {f['domain_name']}",
                    level=2
                )
                doc.add_paragraph(f"Domain compliance: {f.get('domain_compliance', '')}")

            doc.add_heading(f"Penemuan / Finding: {f['id']}", level=3)
            finding_fields = [
                ("ID Penemuan / Finding ID", f["id"]),
                ("Tajuk / Title", f["title"]),
                ("Rujukan CoP / CoP Reference", f["cop_ref"]),
                ("Rujukan Klausa / Clause Reference", f["clause_ref"]),
                ("Sistem (jika berkenaan) / System", f.get("system", "N/A")),
                ("Keterukan / Severity", f["severity"]),
                ("Keadaan Sebenar / Condition", f["condition"]),
                ("Kriteria / Criteria", f["criteria"]),
                ("Punca Masalah / Cause", f["cause"]),
                ("Impak / Consequence", f["consequence"]),
                ("Peluang Penambahbaikan / Corrective Action", f["corrective"]),
                ("Klasifikasi Undang-Undang / Legal Classification", f.get("legal_class", "CoP")),
                ("Boleh Dikompaun? / Compoundable?", f.get("compoundable", "N/A")),
            ]
            add_finding_table(doc, finding_fields)
            doc.add_paragraph("")
    else:
        # Template version — show one example domain and finding template
        doc.add_heading("6.1 CoP 4.0 — Governance and Leadership", level=2)
        doc.add_paragraph("Domain compliance: [N] P / [N] SP / [N] TP / [N] TB out of 1 Tier 1 requirement")

        doc.add_heading("Penemuan / Finding: CoP-2026-001", level=3)
        finding_fields = [
            ("ID Penemuan / Finding ID", "CoP-2026-001"),
            ("Tajuk / Title", "[Concise bilingual title]"),
            ("Rujukan CoP / CoP Reference", "[e.g., CoP 14.3 System and Network Security]"),
            ("Rujukan Klausa / Clause Reference", "[e.g., CoP 14.3.4.1 — specific verbatim clause]"),
            ("Sistem / System (if applicable)", "[NCII system name or N/A for org-level]"),
            ("Keterukan / Severity", "[Kritikal / Tinggi / Sederhana / Rendah]"),
            ("Keadaan Sebenar / Condition", "[What was found — specific evidence, dates, scope, system]"),
            ("Kriteria / Criteria", "[Verbatim CoP clause or Act 854 section that is not met]"),
            ("Punca Masalah / Cause", "[Root cause — why it happened]"),
            ("Impak / Consequence", "[Risk/impact on NCII + penalty exposure if applicable]"),
            ("Peluang Penambahbaikan / Corrective Action", "[Specific, actionable recommendation]"),
            ("Klasifikasi Undang-Undang / Legal Classification", "[Statutory / CE Directive / CoP]"),
            ("Boleh Dikompaun? / Compoundable?", "[Yes — 50% of max fine / No — must prosecute]"),
        ]
        add_finding_table(doc, finding_fields)

        doc.add_paragraph("")
        p = doc.add_paragraph(
            "[Repeat for each CoP domain 4.0-21.0. Include all Tier 1 findings per domain. "
            "Tier 2 findings go in Section 10. If a domain has no findings, state:\n"
            "\"Tiada penemuan ketidakpatuhan bagi domain ini. / "
            "No non-compliance findings identified for this domain.\"]"
        )
        p.italic = True

        doc.add_paragraph("")
        doc.add_heading("Domain-Level Control Counts (Reference)", level=2)
        doc.add_paragraph(
            "The following table shows the number of controls and Tier 1 requirements per domain "
            "to be assessed. Use this to verify completeness of testing."
        ).italic = True
        domain_rows = []
        for section, name, _, controls, t1, t2 in DOMAIN_DATA:
            domain_rows.append((section, name, str(controls), str(t1), str(t2)))
        domain_rows.append(("", "JUMLAH / TOTAL", "126", "210", "18"))
        add_styled_table(doc,
            ["CoP", "Domain", "Kawalan / Controls", "Tier 1", "Tier 2"],
            domain_rows, col_widths=[1.5, 6, 2, 2, 2])


def add_directive_compliance_section(doc, data=None):
    """Section 7: CE Directive Compliance."""
    doc.add_page_break()
    doc.add_heading("7. Pematuhan Arahan Ketua Eksekutif / CE Directive Compliance", level=1)
    doc.add_paragraph(
        "This section covers 42 mandatory requirements extracted from CE Directives 01-10 "
        "that are not duplicated in CoP domains. These are tested separately from CoP-based requirements."
    ).italic = True

    directive_summary = [
        ("CE Directive 01", "CSIRT Establishment", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
        ("CE Directive 02", "Authorised Persons Registration", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
        ("CE Directive 03", "Information Provision to CE", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
        ("CE Directive 04", "NCII System Information", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
        ("CE Directive 05", "Risk Assessment Process", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
        ("CE Directive 06", "Incident Notification", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
        ("CE Directive 07", "Threat Information Sharing", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
        ("CE Directive 08", "Audit Requirements", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
        ("CE Directive 09", "Compliance Monitoring", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
        ("CE Directive 10", "Sector-Specific Directives", "[N]", "[N] P / [N] SP / [N] TP / [N] TB"),
    ]

    if data:
        directive_summary = data.get("directive_summary", directive_summary)

    add_styled_table(doc,
        ["Arahan / Directive", "Perkara / Subject", "Keperluan / Req Count", "Status Pematuhan / Compliance"],
        directive_summary, col_widths=[3, 4, 2, 7])

    doc.add_paragraph("")
    if data and data.get("directive_findings"):
        doc.add_heading("7.1 Penemuan Arahan / Directive Findings", level=2)
        for f in data["directive_findings"]:
            doc.add_heading(f"Penemuan / Finding: {f['id']}", level=3)
            finding_fields = [
                ("ID Penemuan / Finding ID", f["id"]),
                ("Tajuk / Title", f["title"]),
                ("Arahan / Directive", f["directive"]),
                ("Rujukan Perenggan / Paragraph Reference", f["para_ref"]),
                ("Keterukan / Severity", f["severity"]),
                ("Keadaan Sebenar / Condition", f["condition"]),
                ("Kriteria / Criteria", f["criteria"]),
                ("Punca Masalah / Cause", f["cause"]),
                ("Impak / Consequence", f["consequence"]),
                ("Peluang Penambahbaikan / Corrective Action", f["corrective"]),
            ]
            add_finding_table(doc, finding_fields)
            doc.add_paragraph("")
    else:
        doc.add_paragraph(
            "[For each directive with non-compliance findings, include finding details using the same "
            "5 C's framework as Section 6. Directive findings reference the directive paragraph number "
            "instead of a CoP clause number.]"
        ).italic = True


def add_domain_compliance_profile(doc, data=None):
    """Section 8: Domain Compliance Profile."""
    doc.add_page_break()
    doc.add_heading("8. Profil Pematuhan Domain / Domain Compliance Profile", level=1)
    doc.add_paragraph(
        "Aggregation uses weakest-link per domain — domains are presented as a compliance profile, "
        "not blended into a single score. Each domain shows counts for Tier 1 (SHALL) requirements only."
    ).italic = True

    if data:
        domain_rows = data.get("domain_profile", [])
    else:
        domain_rows = []
        for section, name, _, controls, t1, t2 in DOMAIN_DATA:
            domain_rows.append((section, name, str(t1), "", "", "", ""))
        domain_rows.append(("", "JUMLAH / TOTAL", "210", "", "", "", ""))

    add_styled_table(doc,
        ["CoP", "Domain", "Jumlah T1 / T1 Total", "# P", "# SP", "# TP", "# TB"],
        domain_rows, col_widths=[1.3, 5, 2, 1.5, 1.5, 1.5, 1.5])


def add_finding_summary_section(doc, data=None):
    """Section 9: Finding Summary / Register."""
    doc.add_page_break()
    doc.add_heading("9. Ringkasan Penemuan / Finding Summary", level=1)

    doc.add_heading("9.1 Daftar Penemuan / Finding Register", level=2)
    if data:
        register_rows = data.get("register", [])
    else:
        register_rows = [
            ("[CoP-YYYY-NNN]", "[Title]", "[CoP X.Y.Z.N]", "[K/T/S/R]", "[Open]", "[DD/MM/YYYY]"),
        ] * 3
    add_styled_table(doc,
        ["ID", "Tajuk / Title", "Klausa / Clause", "Keterukan / Severity",
         "Status", "Tarikh Sasaran / Target"],
        register_rows, col_widths=[2.5, 4, 2.5, 1.5, 1.5, 2.5])


def add_improvement_section(doc, data=None):
    """Section 10: Improvement Opportunities (Tier 2)."""
    doc.add_page_break()
    doc.add_heading("10. Peluang Penambahbaikan / Improvement Opportunities (Tier 2)", level=1)
    doc.add_paragraph(
        "CE Directive 08 para 23(i) requires improvement opportunities in the report. "
        "These are from Tier 2 (SHOULD/MAY) requirements. They carry NO severity rating, "
        "create NO remediation obligation, and are NOT counted as findings."
    ).italic = True

    if data:
        for opp in data.get("opportunities", []):
            doc.add_heading(f"IO-{opp['num']}: {opp['title']}", level=3)
            fields = [
                ("Rujukan Klausa / Clause Reference", opp["clause_ref"]),
                ("Tahap Obligasi / Obligation Level", opp["obligation"]),
                ("Pemerhatian / Observation", opp["observation"]),
                ("Cadangan / Recommendation", opp["recommendation"]),
            ]
            add_finding_table(doc, fields)
            doc.add_paragraph("")
    else:
        doc.add_heading("IO-001: [Title]", level=3)
        fields = [
            ("Rujukan Klausa / Clause Reference", "[e.g., CoP 5.1.1.2 — SHOULD clause]"),
            ("Tahap Obligasi / Obligation Level", "SHOULD / MAY"),
            ("Pemerhatian / Observation", "[What was observed — gap in SHOULD/MAY requirement]"),
            ("Cadangan / Recommendation", "[Recommended improvement action]"),
        ]
        add_finding_table(doc, fields)
        doc.add_paragraph("")
        doc.add_paragraph(
            "[Repeat for each Tier 2 requirement where an improvement opportunity is identified. "
            "There are 18 Tier 2 requirements across the CoP.]"
        ).italic = True


def add_conclusion_section(doc, data=None):
    """Section 11: Conclusion."""
    doc.add_page_break()
    doc.add_heading("11. Rumusan / Conclusion", level=1)
    doc.add_paragraph(
        "CE Directive 08 para 23(j): \"Gambaran mengenai tahap pematuhan dan keberkesanan "
        "kawalan keselamatan siber yang dilaksanakan oleh entiti IKMN. Rumusan perlu menekankan "
        "aspek utama yang memerlukan perhatian dan tindakan segera.\""
    ).italic = True

    doc.add_heading("11.1 Ringkasan Postur Pematuhan / Compliance Posture Summary", level=2)
    if data:
        for item in data.get("posture", []):
            doc.add_paragraph(item)
    else:
        doc.add_paragraph("[Total Tier 1 requirements assessed, P/SP/TP/TB counts and percentages]")
        doc.add_paragraph("[Directive requirements compliance summary]")
        doc.add_paragraph("[Domains with strongest compliance]")
        doc.add_paragraph("[Domains with weakest compliance]")

    doc.add_heading("11.2 Ringkasan Penilaian Risiko / Risk Assessment Summary", level=2)
    if data:
        for item in data.get("risk_summary", []):
            doc.add_paragraph(item)
    else:
        doc.add_paragraph("[Number and distribution of findings by severity]")
        doc.add_paragraph("[Key risk areas — greatest threat to NCII]")
        doc.add_paragraph("[Relationship between non-compliance and risk to national critical infrastructure]")

    doc.add_heading("11.3 Perkara Memerlukan Perhatian / Key Areas Requiring Attention", level=2)
    if data:
        for item in data.get("attention", []):
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("[Critical and High findings requiring immediate remediation]")
        doc.add_paragraph("[Systemic patterns across multiple domains]")

    doc.add_heading("11.4 Susulan Audit Terdahulu / Prior Audit Follow-Up", level=2)
    if data:
        doc.add_paragraph(data.get("prior_audit", ""))
    else:
        doc.add_paragraph("[Status of prior audit findings: closed / open / recurring]")
        doc.add_paragraph("[Whether recurring findings indicate systemic compliance failure]")

    doc.add_heading("11.5 Pendapat Tambahan / Supplementary Opinion", level=2)
    if data:
        opinion_rows = data.get("opinion_rows", [])
    else:
        opinion_rows = [
            ("Memuaskan (Satisfactory)",
             "No Critical/High findings. TP count zero or near-zero.", "[ ]"),
            ("Perlu Penambahbaikan (Needs Improvement)",
             "One or more High, OR multiple Medium with systemic gaps.", "[ ]"),
            ("Tidak Memuaskan (Unsatisfactory)",
             "One or more Critical, OR multiple High across core domains.", "[ ]"),
        ]
    add_styled_table(doc,
        ["Pendapat / Opinion", "Kriteria / Criteria", "Dipilih / Selected"],
        opinion_rows, col_widths=[4, 9, 2])

    doc.add_paragraph("")
    doc.add_heading("11.6 Input Pengamal Undang-Undang / Legal Practitioner Input (Mode A)", level=2)
    if data:
        for item in data.get("legal_input", []):
            doc.add_paragraph(item)
    else:
        doc.add_paragraph("[Statutory compliance conclusions reviewed by legal practitioner for legal accuracy]")
        doc.add_paragraph("[Penalty exposure for Critical and High findings]")
        doc.add_paragraph("[Note: Mode B privileged opinion is NOT referenced in this report]")


def add_map_section(doc, data=None):
    """Section 12: Management Action Plan."""
    doc.add_page_break()
    doc.add_heading("12. Pelan Tindakan Pengurusan / Management Action Plan (MAP)", level=1)
    doc.add_paragraph(
        "Management's response to each finding, including agreed remediation actions, "
        "responsible owners, and target completion dates. Per CE Directive 08 para 11, "
        "the entity must confirm findings before report finalisation."
    ).italic = True

    if data:
        map_rows = data.get("map_rows", [])
    else:
        map_rows = [
            ("[CoP-YYYY-NNN]", "[Agree/Disagree/Partial]", "[Specific action]",
             "[Name]", "[DD/MM/YYYY]", "[Open]"),
        ] * 3
    add_styled_table(doc,
        ["ID Penemuan / Finding ID", "Respons / Response", "Tindakan / Action",
         "Pemilik / Owner", "Tarikh / Target", "Status"],
        map_rows, col_widths=[2, 2.5, 4, 2, 2, 1.5])


def add_appendices(doc, data=None):
    """Appendices A-D."""
    # Lampiran A — NCII System Register
    doc.add_page_break()
    doc.add_heading("Lampiran A — Senarai Sistem NCII / NCII System Register", level=1)
    if data:
        system_rows = data.get("systems", [])
    else:
        system_rows = [
            ("[NCII-SYS-001]", "[Name]", "[Description]", "[Critical/Standard]", "[Location]", "[Owner]"),
        ] * 3
    add_styled_table(doc,
        ["ID Sistem / System ID", "Nama / Name", "Perihalan / Description",
         "Klasifikasi / Classification", "Lokasi / Location", "Pemilik / Owner"],
        system_rows, col_widths=[2, 3, 3.5, 2, 2, 2.5])

    # Lampiran B — Penalty Exposure
    doc.add_page_break()
    doc.add_heading("Lampiran B — Pendedahan Penalti / Penalty Exposure Summary", level=1)
    doc.add_paragraph(
        "Note: This appendix summarises penalty exposure based on audit findings (Mode A input). "
        "The full privileged penalty quantification (Mode B) is a separate, confidential deliverable "
        "to the audit committee and is NOT included in this report."
    ).italic = True

    if data:
        penalty_rows = data.get("penalties", [])
    else:
        penalty_rows = [
            ("[s22]", "[Offence description]", "[RM X]", "[RM X]", "[Yes/No]",
             "[Finding ID]"),
        ] * 3
    add_styled_table(doc,
        ["Seksyen / Section", "Kesalahan / Offence", "Denda Maks / Max Fine",
         "Penjara Maks / Max Imprisonment", "Boleh Dikompaun / Compoundable",
         "Penemuan / Finding"],
        penalty_rows, col_widths=[1.5, 4, 2, 2.5, 2, 2.5])

    # Lampiran C — Severity Definitions
    doc.add_page_break()
    doc.add_heading("Lampiran C — Peringkat Keterukan / Severity Definitions", level=1)
    severity_defs = [
        ("Kritikal\n(Critical)",
         "Direct and immediate threat to NCII. Exploitation likely or occurring. "
         "National-level impact on essential services. Criminal liability exposure under Act 854. "
         "Triggers s52 director liability assessment. Remediation: 30 days.",
         "30 hari / 30 days"),
        ("Tinggi\n(High)",
         "Significant vulnerability that could compromise NCII. Material impact on "
         "essential services. Non-compliance with statutory obligation (s17-s27) or critical CoP control. "
         "Remediation: 90 days.",
         "90 hari / 90 days"),
        ("Sederhana\n(Medium)",
         "Moderate risk — exploitation requires additional factors. Compensating "
         "controls partially mitigate. Non-compliance with CoP requirement. "
         "Remediation: 180 days.",
         "180 hari / 180 days"),
        ("Rendah\n(Low)",
         "Minor risk — theoretical vulnerability with low likelihood. Limited impact. "
         "Documentation or process gap. Remediation: next audit cycle.",
         "Kitaran audit seterusnya / Next audit cycle"),
    ]
    add_styled_table(doc,
        ["Keterukan / Severity", "Definisi / Definition", "Sasaran Pemulihan / Remediation Target"],
        severity_defs, col_widths=[2.5, 10, 3.5])

    # Lampiran D — Glossary
    doc.add_page_break()
    doc.add_heading("Lampiran D — Glosari / Glossary", level=1)
    glossary = [
        ("Act 854", "Cyber Security Act 2024 / Akta Keselamatan Siber 2024"),
        ("BIA", "Business Impact Analysis"),
        ("CAB", "Change Advisory Board"),
        ("CE", "Chief Executive, NACSA / Ketua Eksekutif, NACSA"),
        ("CoP", "Code of Practice / Tataamalan"),
        ("CSIMP", "Cyber Security Incident Management Plan / Prosedur Pengurusan Insiden Keselamatan Siber"),
        ("CSIRT", "Cyber Security Incident Response Team"),
        ("CSSP", "Cyber Security Service Provider"),
        ("DLP", "Data Leakage Prevention"),
        ("IKMN", "Infrastruktur Maklumat Kritikal Negara (NCII)"),
        ("MAP", "Management Action Plan / Pelan Tindakan Pengurusan"),
        ("MTTD", "Mean Time to Detect"),
        ("MTTR", "Mean Time to Respond"),
        ("NC4", "National Cyber Coordination and Command Centre"),
        ("NCII", "National Critical Information Infrastructure / IKMN"),
        ("NCSB", "National Cyber Security Baseline"),
        ("P", "Patuh / Compliant"),
        ("PAM", "Privileged Access Management"),
        ("P.U.(A)", "Subsidiary legislation / Perundangan Undang-Undang"),
        ("Rumusan", "Conclusion / Summary"),
        ("SMP", "Security Monitoring Procedure"),
        ("SP", "Separa Patuh / Partially Compliant"),
        ("TB", "Tidak Berkenaan / Not Applicable"),
        ("TP", "Tidak Patuh / Non-Compliant"),
        ("VMP", "Vulnerability Management Programme"),
    ]
    add_styled_table(doc, ["Istilah / Term", "Definisi / Definition"],
                     glossary, col_widths=[3, 13])


def add_signoff(doc, data=None):
    """Sign-off page."""
    doc.add_page_break()
    doc.add_heading("Tandatangan / Sign-Off", level=1)

    if data:
        sign_offs = data.get("signoffs", [])
    else:
        sign_offs = [
            ("Disediakan oleh / Prepared by:", "Ketua Juruaudit / Lead Auditor",
             "[Name]", "[Date]"),
            ("Disemak oleh / Reviewed by:", "Rakan Kongsi Jaminan / Engagement Partner",
             "[Name]", "[Date]"),
            ("Semakan undang-undang / Legal review (Mode A):", "Pengamal Undang-Undang / Legal Practitioner",
             "[Name]", "[Date]"),
            ("Pengesahan entiti / Entity confirmation:", "Pegawai diberi kuasa / Authorised Officer",
             "[Name]", "[Date]"),
            ("Dihantar kepada NACSA / Submitted to NACSA:", "Via auditpematuhan@nacsa.gov.my",
             "Salinan / Copy #", "[Date]"),
        ]

    for label, role, name, d in sign_offs:
        p = doc.add_paragraph()
        run = p.add_run(f"{label} ")
        run.bold = True
        p.add_run(f"{role}\n")
        p.add_run(f"Nama / Name: {name}\n")
        p.add_run(f"Tarikh / Date: {d}\n")
        p.add_run("Tandatangan / Signature: _________________________")
        doc.add_paragraph("")

    # Footer note
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        "Laporan ini hendaklah dihantar kepada Ketua Eksekutif NACSA dalam tempoh 30 hari "
        "selepas audit selesai melalui auditpematuhan@nacsa.gov.my menurut s23 dan "
        "Arahan Ketua Eksekutif 08. Tempoh penyimpanan: 7 tahun.\n"
        "This report must be submitted to the Chief Executive of NACSA within 30 days of audit "
        "completion via auditpematuhan@nacsa.gov.my per s23 and CE Directive 08. "
        "Retention period: 7 years."
    )
    run.italic = True
    run.font.size = Pt(8)


# ============================================================================
# DOCUMENT 1: REPORT TEMPLATE
# ============================================================================

def build_report_template(output_path):
    """Generate the blank report template."""
    doc = Document()
    setup_styles(doc)

    entity_data = {}  # Empty = template placeholders
    add_cover_page(doc, entity_data)
    add_distribution(doc, is_sample=False)
    add_toc(doc)
    add_executive_summary(doc)
    add_scope_section(doc)
    add_methodology_section(doc)
    add_statutory_section(doc)
    add_director_liability_section(doc)
    add_findings_section(doc)
    add_directive_compliance_section(doc)
    add_domain_compliance_profile(doc)
    add_finding_summary_section(doc)
    add_improvement_section(doc)
    add_conclusion_section(doc)
    add_map_section(doc)
    add_appendices(doc)
    add_signoff(doc)

    doc.save(output_path)
    print(f"Template saved: {output_path}")


# ============================================================================
# DOCUMENT 2: SAMPLE AUDIT REPORT
# ============================================================================

def build_sample_report(output_path):
    """Generate the fictional sample audit report for Syarikat Tenaga Siber Sdn Bhd."""
    doc = Document()
    setup_styles(doc)

    # --- Entity data ---
    entity_data = {
        "entity_name": "Syarikat Tenaga Siber Sdn Bhd",
        "sector": "Tenaga / Energy",
        "sector_lead": "Suruhanjaya Tenaga / Energy Commission (ST)",
        "system_count": "4 sistem NCII",
        "period": "06/01/2026 hingga 14/02/2026",
        "report_date": "24/02/2026",
        "audit_firm": "CyberAssurance Partners Sdn Bhd",
        "partner": "Dato' Ir. Razlan Hashim, CISA, CISSP, PMP",
        "lead_auditor": "Ahmad Faizal bin Mohd Noor, CISA, CISM, ISO 27001 LA",
        "legal": "Puan Siti Nurhaliza binti Ahmad, Malaysian Bar — Azman Rahim & Partners",
        "ref_no": "NACSA-AUDIT-2026-0147",
    }

    add_cover_page(doc, entity_data)
    add_distribution(doc, is_sample=True)
    add_toc(doc)

    # --- Section 1: Executive Summary ---
    exec_data = {
        "intro": (
            "Syarikat Tenaga Siber Sdn Bhd (\"STS\" atau \"Entiti\") merupakan sebuah entiti NCII sektor tenaga "
            "yang beroperasi di bawah Suruhanjaya Tenaga sebagai ketua sektor. STS memiliki dan mengawal "
            "4 sistem NCII yang ditetapkan: sistem kawalan SCADA, sistem pengurusan tenaga (EMS), platform "
            "grid pintar, dan rangkaian korporat. Ini merupakan audit keselamatan siber pertama bagi STS "
            "sejak pelantikannya sebagai entiti NCII.\n\n"
            "STS is a mid-tier energy sector NCII entity operating under the Energy Commission as sector lead. "
            "STS owns and controls 4 designated NCII systems: SCADA control system, energy management system (EMS), "
            "smart grid platform, and corporate network. This is the first cyber security audit since STS's "
            "designation as an NCII entity."
        ),
        "metrics": [
            ("Total Tier 1 CoP requirements assessed", "210 of 210"),
            ("Directive requirements assessed", "42 of 42"),
            ("Patuh (Compliant)", "168 (80.0%)"),
            ("Separa Patuh (Partially Compliant)", "24 (11.4%)"),
            ("Tidak Patuh (Non-Compliant)", "14 (6.7%)"),
            ("Tidak Berkenaan (Not Applicable)", "4 (1.9%)"),
            ("Tier 2 improvement opportunities identified", "3 of 18"),
            ("Pendapat Tambahan / Supplementary Opinion", "Perlu Penambahbaikan (Needs Improvement)"),
        ],
        "findings_summary": [
            ("Kritikal (Critical)", "1", "CSIRT not established; 6-hour notification untested (CoP 18.0)"),
            ("Tinggi (High)", "2", "Risk assessment submission overdue; SCADA firewall rules stale (CoP 9.0, 14.0)"),
            ("Sederhana (Medium)", "2", "Partial vulnerability scanning; DLP gap on SCADA (CoP 16.0, 11.0)"),
            ("Rendah (Low)", "1", "Physical security procedure review overdue (CoP 13.0)"),
            ("Jumlah / Total", "6", ""),
        ],
        "immediate_attention": [
            "KRITIKAL: Establishment of a CSIRT and testing of the 6-hour incident notification process "
            "is the most urgent priority. Without a functioning CSIRT and tested notification process, "
            "STS cannot comply with its s26 statutory obligation and is exposed to criminal liability "
            "under Act 854.",
            "TINGGI: The risk assessment report must be submitted to the NACSA CE via "
            "auditpematuhan@nacsa.gov.my within the prescribed period. Current overdue status of "
            "52 days exposes the entity to penalties under s21.",
            "TINGGI: SCADA firewall rules have not been reviewed for 18 months — well beyond the "
            "annual review mandate under CoP 14.3.4.1. Given the criticality of the SCADA system "
            "to national energy infrastructure, this represents a significant risk.",
        ],
    }
    add_executive_summary(doc, exec_data)

    # --- Section 2: Scope ---
    scope_data = {
        "scope_rows": [
            ("NCII Systems in Scope", "4 systems — see Lampiran A"),
            ("Tier 1 CoP Requirements", "210 (SHALL clauses from 126 controls across 17 active domains)"),
            ("Tier 2 CoP Requirements", "18 (SHOULD/MAY clauses — improvement opportunities only)"),
            ("CE Directive Requirements", "42 (from CE Directives 01-10)"),
            ("Total Testable Requirements", "210 Tier 1 + 42 directive = 252 compliance-rated"),
            ("CoP Domains Assessed", "4.0-20.0 (17 active domains; 21.0 assessed but 0 requirements)"),
            ("Assessment Period", "06/01/2026 to 14/02/2026 (6 weeks)"),
            ("Exclusions", "CoP 20.0 sector-specific: Energy Commission sector CoP not yet gazetted — "
             "tested against CoP Template requirements only"),
            ("Scope Limitations", "Smart grid platform underwent maintenance during Week 3; "
             "technical testing of CoP 14.0 controls for this system deferred to Week 5"),
        ],
        "team": [
            ("Engagement Partner", "Dato' Ir. Razlan Hashim", "CISA, CISSP, PMP",
             "Overall direction, QA, NACSA submission"),
            ("Lead Auditor (Ketua Juruaudit)", "Ahmad Faizal bin Mohd Noor", "CISA, CISM, ISO 27001 LA",
             "Full CoP scope (4.0-21.0), directive compliance"),
            ("Auditor (Juruaudit)", "Nurul Aisyah binti Zainal", "CISA, CEH",
             "CoP 14.0-16.0 (network, access, vulnerability)"),
            ("Auditor (Juruaudit)", "Muhammad Hafiz bin Ismail", "ISO 27001 LA",
             "CoP 4.0-10.0 (governance, policy, risk)"),
            ("Legal Practitioner (Mode A)", "Puan Siti Nurhaliza binti Ahmad", "Malaysian Bar",
             "Statutory obligations s17-s27, s52 director liability, penalty exposure"),
            ("Sector Specialist", "Ir. Lee Wei Ming", "P.Eng, GICSP",
             "SCADA/ICS security, energy sector controls, CoP 20.0"),
        ],
    }
    add_scope_section(doc, scope_data)

    # --- Section 3: Methodology ---
    method_data = {
        "optional_approaches": [
            ("Control-Based (Kawalan)", "Not Elected", "N/A"),
            ("Technical Testing (Ujian Teknikal)", "Not Elected", "N/A"),
            ("Inspection & Verification (Pemeriksaan)", "Not Elected", "N/A"),
            ("Continuous Improvement (Penambahbaikan)", "Not Elected", "N/A"),
        ],
        "sampling": (
            "Judgmental (non-statistical) sampling applied for all 4 NCII systems. "
            "Org-level controls (145 requirements) tested once. System-level controls (65 requirements) "
            "tested per NCII system where applicable. Total test procedures executed: 405. "
            "Document review, interviews, observation, and system inspection performed."
        ),
    }
    add_methodology_section(doc, method_data)

    # --- Section 4: Statutory Obligations ---
    stat_data = {
        "obligations": [
            ("s17", "Designation as NCII Entity\nPerlantikan sebagai Entiti IKMN",
             "Met", "STS designated 15/03/2025 via written notice from CE."),
            ("s18", "General Duties of NCII Entity\nKewajipan Am Entiti IKMN",
             "Met", "CISO appointed. Board-level cyber security oversight established."),
            ("s19", "Notification of Change of Ownership\nPemberitahuan Perubahan Pemilikan",
             "Met", "No ownership changes since designation. Process documented."),
            ("s20", "Provision of Information\nPemberian Maklumat",
             "Met", "Entity responded to NACSA information requests within prescribed periods."),
            ("s21", "Risk Assessment\nPenilaian Risiko",
             "NOT MET", "Risk assessment conducted but report NOT submitted to CE within 30 days. "
             "Report overdue by 52 days as at audit date. See Finding CoP-2026-002."),
            ("s22", "Audit and CoP Compliance\nAudit dan Pematuhan Tataamalan",
             "Met", "This audit conducted within prescribed period. CoP compliance assessed herein."),
            ("s23", "Audit Report Submission\nPenghantaran Laporan Audit",
             "Pending", "Report to be submitted within 30 days of audit completion."),
            ("s24", "Compliance with Standards\nPematuhan Piawaian",
             "Met", "No additional standards imposed by CE beyond CoP."),
            ("s25", "Cyber Threat Information Sharing\nPerkongsian Maklumat Ancaman",
             "NOT MET", "Entity not registered as authorised person with NACSA for threat intel sharing. "
             "See Finding DIR-2026-001."),
            ("s26", "Incident Notification (6-hour)\nPemberitahuan Insiden (6 jam)",
             "NOT MET", "CSIRT not established. 6-hour notification process not tested. "
             "NC4 portal access not configured. See Finding CoP-2026-001."),
            ("s27", "Compliance with CE Directives\nPematuhan Arahan KE",
             "Partial", "Non-compliance with Directive 01 (CSIRT) and Directive 05 (RA submission)."),
        ],
    }
    add_statutory_section(doc, stat_data)

    # --- Section 5: Director Liability ---
    director_data = {
        "assessment": [
            "Three statutory obligations assessed as NOT MET: s21 (risk assessment submission), "
            "s25 (threat information sharing), and s26 (incident notification). Per s52 of Act 854, "
            "these non-compliances constitute offences that may trigger director liability.",
            "The audit team found that management was aware of the CSIRT requirement (s26/Directive 01) "
            "but had not allocated budget for its establishment. The Board approved the FY2026 budget "
            "without a CSIRT line item despite the CISO's recommendation.",
            "For s21, the CISO completed the risk assessment on time but the report was pending "
            "CEO sign-off for 52 days due to competing priorities. This represents a failure of "
            "management oversight, not a technical capability gap.",
            "Due diligence defence under s52(2): The CISO can demonstrate that the offence was "
            "committed without their consent (they recommended the CSIRT budget and flagged the "
            "s21 submission deadline). The CEO and CFO face higher s52 exposure.",
        ],
        "officers": [
            ("Chief Executive Officer", "Encik Kamal bin Yusoff",
             "High", "Approved budget without CSIRT allocation despite CISO recommendation"),
            ("Chief Financial Officer", "Puan Farah binti Hassan",
             "High", "Did not approve CSIRT budget request; delayed s21 report sign-off"),
            ("Chief Information Security Officer", "Encik Rizwan bin Abdullah",
             "Low", "Documented recommendations for CSIRT and s21 compliance; overruled by management"),
            ("Chief Technology Officer", "Ir. Tan Mei Ling",
             "Medium", "Did not escalate SCADA firewall review gap; responsible for technical controls"),
        ],
    }
    add_director_liability_section(doc, director_data)

    # --- Section 6: Findings by Domain (Tier 1) ---
    findings = [
        {
            "domain": "incident-management",
            "domain_section": "18.0",
            "domain_name": "Cyber Security Incident Management",
            "domain_compliance": "2 P / 1 SP / 3 TP / 0 TB out of 6 Tier 1 requirements",
            "id": "CoP-2026-001",
            "title": "CSIRT Tidak Ditubuhkan, Proses Pemberitahuan 6 Jam Tidak Diuji / "
                     "CSIRT Not Established, 6-Hour Notification Process Untested",
            "cop_ref": "CoP 18.0 — Cyber Security Incident Management",
            "clause_ref": "CoP 18.1.1.1 — \"The NCII entity shall establish a Cyber Security Incident "
                         "Management Procedure (CSIMP) to ensure effective preparation and response to "
                         "cyber security incidents.\"",
            "system": "All 4 NCII systems (org-level control)",
            "severity": "Kritikal (Critical)",
            "condition": (
                "Audit inspection confirmed that STS has not established a Cyber Security Incident "
                "Response Team (CSIRT) as at the audit date (14/02/2026). No CSIMP document exists. "
                "The entity has no defined incident escalation procedure, no designated incident "
                "handlers, and no tested process for notifying the NACSA CE within the mandatory "
                "6-hour window via the NC4 portal. The NC4 portal account has not been activated. "
                "Interview with the CISO (Encik Rizwan, 22/01/2026) confirmed that a CSIRT "
                "establishment proposal was submitted to the Board in Q3 2025 but was not approved "
                "due to budget constraints."
            ),
            "criteria": (
                "CoP 18.1.1.1 requires establishment of a CSIMP. "
                "CE Directive 01 requires CSIRT establishment. "
                "s26 of Act 854 requires notification of cyber security incidents to the CE "
                "within 6 hours via NC4 portal per P.U.(A) 220/2024. "
                "CE Directive 06 prescribes the notification process."
            ),
            "cause": (
                "Management did not allocate budget for CSIRT establishment in FY2025 or FY2026 "
                "despite the CISO's recommendation. The Board prioritised capital expenditure on "
                "smart grid expansion over cyber security capability building."
            ),
            "consequence": (
                "Without a CSIRT, the entity cannot detect, classify, or respond to cyber security "
                "incidents affecting the SCADA control system or other NCII. A ransomware attack or "
                "state-sponsored intrusion targeting the energy grid could proceed undetected and "
                "unreported, potentially disrupting electricity supply to ~180,000 consumers in the "
                "Northern Region. Non-compliance with s26 exposes the entity to a fine not exceeding "
                "RM500,000 or imprisonment not exceeding 10 years or both. s52 director liability "
                "applies."
            ),
            "corrective": (
                "1. Establish CSIRT with defined roles, responsibilities, and authority within 30 days.\n"
                "2. Develop and approve CSIMP covering all 4 NCII systems.\n"
                "3. Activate NC4 portal access and configure 6-hour notification workflow.\n"
                "4. Conduct tabletop exercise simulating a SCADA-targeted incident.\n"
                "5. Register authorised persons with NACSA per CE Directive 02."
            ),
            "legal_class": "Statutory (s26) + CE Directive 01 + CoP 18.1.1",
            "compoundable": "s26: No — not a compoundable offence. Must prosecute if charged.",
        },
        {
            "domain": "risk-management",
            "domain_section": "9.0",
            "domain_name": "Risk Management",
            "domain_compliance": "14 P / 3 SP / 2 TP / 0 TB out of 19 Tier 1 requirements",
            "id": "CoP-2026-002",
            "title": "Laporan Penilaian Risiko Tidak Dihantar Dalam Tempoh / "
                     "Risk Assessment Report Not Submitted Within Prescribed Period",
            "cop_ref": "CoP 9.0 — Risk Management",
            "clause_ref": "CoP 9.1.4.2 — \"The report shall be submitted to the Chief Executive of NACSA "
                         "via e-mail of auditpematuhan@nacsa.gov.my, within 30 days upon the completion "
                         "of the risk assessment.\"",
            "system": "N/A (org-level control)",
            "severity": "Tinggi (High)",
            "condition": (
                "STS completed its cyber security risk assessment on 15/10/2025. The risk assessment "
                "report was finalised internally on 28/10/2025. However, the report had not been "
                "submitted to the NACSA CE as at the audit date (14/02/2026) — 112 days after "
                "completion and 82 days past the 30-day submission deadline. The CISO confirmed "
                "(interview, 23/01/2026) that the report was pending CEO sign-off. Review of the "
                "risk assessment report (Doc Ref: STS-RA-2025-001) confirmed it was substantively "
                "complete and aligned with CE Directive 05 requirements."
            ),
            "criteria": (
                "CoP 9.1.4.2 requires submission to NACSA CE within 30 days of completion. "
                "CE Directive 05 para 8.5 prescribes submission via auditpematuhan@nacsa.gov.my. "
                "s21 of Act 854 mandates the risk assessment obligation."
            ),
            "cause": (
                "Internal approval process requires CEO sign-off on all NACSA submissions. "
                "The CEO was occupied with the smart grid platform launch (Q4 2025) and did not "
                "prioritise the risk assessment report for sign-off. No escalation mechanism existed "
                "for time-sensitive regulatory submissions."
            ),
            "consequence": (
                "Late submission of the risk assessment report constitutes non-compliance with s21 "
                "of Act 854, punishable by a fine not exceeding RM200,000 or imprisonment not "
                "exceeding 3 years or both. Additionally, without a submitted risk assessment, "
                "NACSA cannot assess the entity's risk posture, potentially delaying sector-level "
                "risk intelligence. s52 director liability applies."
            ),
            "corrective": (
                "1. Obtain CEO sign-off and submit the risk assessment report to "
                "auditpematuhan@nacsa.gov.my immediately (within 7 days).\n"
                "2. Establish a regulatory submission tracker with automated deadline alerts.\n"
                "3. Delegate submission authority to the CISO for time-sensitive regulatory filings.\n"
                "4. Implement Board reporting on regulatory compliance timelines."
            ),
            "legal_class": "Statutory (s21) + CE Directive 05 + CoP 9.1.4",
            "compoundable": "s21: Yes — compoundable at 50% of maximum fine (RM100,000).",
        },
        {
            "domain": "system-network-security",
            "domain_section": "14.0",
            "domain_name": "System and Network Security",
            "domain_compliance": "18 P / 4 SP / 2 TP / 0 TB out of 24 Tier 1 requirements",
            "id": "CoP-2026-003",
            "title": "Peraturan Firewall SCADA Tidak Disemak Dalam 18 Bulan / "
                     "SCADA Firewall Rules Not Reviewed in 18 Months",
            "cop_ref": "CoP 14.0 — System and Network Security",
            "clause_ref": "CoP 14.3.4.1 — \"Firewall rules, network threat detection capability, and NAC "
                         "solutions shall be reviewed annually to ensure their continued effectiveness in "
                         "detecting and preventing unauthorised access and cyber security threats.\"",
            "system": "NCII-SYS-001: SCADA Control System",
            "severity": "Tinggi (High)",
            "condition": (
                "Review of the SCADA network firewall (Palo Alto PA-5260, FW-SCADA-01) configuration "
                "audit trail confirmed the last firewall rule review was conducted on 12/08/2024 — "
                "18 months before the audit date. The firewall currently has 247 active rules, "
                "including 23 rules marked as 'temporary' dating from 2023 that have not been "
                "reviewed or removed. The entity's own Network Security Policy (NSP-2024-003) "
                "requires quarterly firewall rule reviews for SCADA segments. The ICS network "
                "segment contains 142 SCADA endpoints controlling power distribution to the "
                "Northern Region grid."
            ),
            "criteria": (
                "CoP 14.3.4.1 mandates annual review of firewall rules. "
                "Entity's own NSP-2024-003 requires quarterly review for SCADA segments (more "
                "stringent than CoP, but audited against CoP minimum). "
                "s24 of Act 854 requires compliance with applicable standards."
            ),
            "cause": (
                "The network security team (2 staff) was reassigned to the smart grid platform "
                "deployment project in Q1 2025. No replacement resource was assigned to maintain "
                "the SCADA firewall review schedule. The team lead departed in June 2025 and the "
                "position remained vacant for 4 months."
            ),
            "consequence": (
                "Stale firewall rules on the SCADA segment create an elevated risk of unauthorised "
                "access to industrial control systems. The 23 unreviewed 'temporary' rules may permit "
                "traffic paths that are no longer required, expanding the attack surface. A threat actor "
                "exploiting a stale rule could gain access to SCADA endpoints controlling power "
                "distribution, potentially causing load-shedding or equipment damage affecting "
                "~180,000 consumers."
            ),
            "corrective": (
                "1. Conduct an immediate firewall rule review for FW-SCADA-01 (within 30 days).\n"
                "2. Remove or justify all 23 'temporary' rules from 2023.\n"
                "3. Re-establish quarterly review schedule per entity's own NSP-2024-003.\n"
                "4. Assign a dedicated network security resource for SCADA/OT environments.\n"
                "5. Implement firewall rule lifecycle management with automatic expiry for "
                "temporary rules."
            ),
            "legal_class": "CoP 14.3.4",
            "compoundable": "N/A — CoP non-compliance, not directly a statutory offence.",
        },
        {
            "domain": "technical-vulnerability",
            "domain_section": "16.0",
            "domain_name": "Technical Vulnerability",
            "domain_compliance": "7 P / 2 SP / 1 TP / 0 TB out of 10 Tier 1 requirements",
            "id": "CoP-2026-004",
            "title": "Pengimbasan Kerentanan Hanya Meliputi 70% Aset NCII / "
                     "Vulnerability Scanning Covers Only 70% of NCII Assets",
            "cop_ref": "CoP 16.0 — Technical Vulnerability",
            "clause_ref": "CoP 16.2.1.1 — \"The NCII entity shall define a Vulnerability Assessment Plan "
                         "(VAP) that outline the following: The scope of the assessment, covering all NCII "
                         "assets and associated systems.\"",
            "system": "NCII-SYS-001 (SCADA), NCII-SYS-003 (Smart Grid)",
            "severity": "Sederhana (Medium)",
            "condition": (
                "Review of the Vulnerability Assessment Plan (STS-VAP-2025-001) and Qualys scan "
                "logs confirmed that vulnerability scanning covers the corporate network (NCII-SYS-004) "
                "and the EMS (NCII-SYS-002) comprehensively. However, the SCADA control system "
                "(NCII-SYS-001) and smart grid platform (NCII-SYS-003) are only partially scanned — "
                "approximately 70% of total NCII assets are in scope. The sector specialist confirmed "
                "(interview with Ir. Lee, 05/02/2026) that 43 SCADA endpoints and 18 smart grid "
                "controllers were excluded from scanning due to vendor restrictions on active scanning "
                "of OT devices."
            ),
            "criteria": (
                "CoP 16.2.1.1 requires the VAP to cover 'all NCII assets and associated systems.' "
                "The current plan excludes 30% of assets."
            ),
            "cause": (
                "OT device vendors (Siemens, ABB) prohibit active vulnerability scanning on certain "
                "legacy PLCs and RTUs, citing risk of device malfunction. The entity accepted this "
                "restriction without implementing compensating controls (e.g., passive network scanning, "
                "firmware version tracking, or vendor patch advisory monitoring)."
            ),
            "consequence": (
                "Unscanned OT assets may harbour known vulnerabilities that remain undetected and "
                "unpatched. For SCADA/ICS environments, these vulnerabilities could be exploited by "
                "threat actors targeting energy infrastructure (e.g., Industroyer-class malware). "
                "The risk is partially mitigated by network segmentation between IT and OT segments."
            ),
            "corrective": (
                "1. Implement passive vulnerability scanning for OT assets that cannot tolerate "
                "active scanning.\n"
                "2. Establish firmware version tracking and vendor security advisory monitoring "
                "for all 43 SCADA endpoints and 18 smart grid controllers.\n"
                "3. Update the VAP to document compensating controls for excluded assets.\n"
                "4. Engage OT security vendor to conduct safe OT vulnerability assessment."
            ),
            "legal_class": "CoP 16.2.1",
            "compoundable": "N/A — CoP non-compliance.",
        },
        {
            "domain": "data-security",
            "domain_section": "11.0",
            "domain_name": "Data Security",
            "domain_compliance": "19 P / 3 SP / 1 TP / 0 TB out of 23 Tier 1 requirements",
            "id": "CoP-2026-005",
            "title": "Penyelesaian DLP Tidak Dilaksanakan pada Segmen SCADA / "
                     "DLP Solution Not Deployed on SCADA Segment",
            "cop_ref": "CoP 11.0 — Data Security",
            "clause_ref": "CoP 11.1.1.1 — \"The NCII entity shall implement a DLP solution across its "
                         "organisation, to safeguard sensitive data from unauthorised access, transmission "
                         "or leakage.\"",
            "system": "NCII-SYS-001: SCADA Control System",
            "severity": "Sederhana (Medium)",
            "condition": (
                "STS has deployed Symantec DLP on the corporate network (NCII-SYS-004) and the EMS "
                "(NCII-SYS-002) covering email, web, and USB channels. However, the SCADA control "
                "system (NCII-SYS-001) and smart grid platform (NCII-SYS-003) do not have DLP "
                "coverage. The CISO confirmed (interview, 28/01/2026) that DLP agents were not "
                "installed on SCADA HMI workstations due to concerns about performance impact on "
                "real-time control operations."
            ),
            "criteria": (
                "CoP 11.1.1.1 requires DLP implementation 'across its organisation.' "
                "The current deployment excludes SCADA and smart grid environments."
            ),
            "cause": (
                "Technical concern that DLP agents on SCADA HMI workstations could introduce "
                "latency in real-time control operations. No evaluation of OT-compatible DLP "
                "solutions was conducted."
            ),
            "consequence": (
                "Sensitive SCADA configuration data, PLC programming files, and grid topology "
                "information on HMI workstations are not protected against exfiltration. A malicious "
                "insider or compromised workstation could exfiltrate operational data without "
                "detection. Risk is partially mitigated by network segmentation and USB port "
                "disablement on SCADA HMI stations."
            ),
            "corrective": (
                "1. Evaluate OT-compatible DLP solutions (e.g., network-based DLP that monitors "
                "traffic without endpoint agents).\n"
                "2. Implement network DLP on the SCADA/IT boundary to monitor data flows.\n"
                "3. Document compensating controls for SCADA endpoints where agent-based DLP "
                "is not feasible.\n"
                "4. Include SCADA DLP coverage in the FY2027 cyber security budget."
            ),
            "legal_class": "CoP 11.1.1",
            "compoundable": "N/A — CoP non-compliance.",
        },
        {
            "domain": "physical-security",
            "domain_section": "13.0",
            "domain_name": "Physical Security",
            "domain_compliance": "13 P / 1 SP / 1 TP / 0 TB out of 15 Tier 1 requirements",
            "id": "CoP-2026-006",
            "title": "Prosedur Keselamatan Fizikal Lewat Disemak 3 Bulan / "
                     "Physical Security Procedure Overdue for Review by 3 Months",
            "cop_ref": "CoP 13.0 — Physical Security",
            "clause_ref": "CoP 13.1.5.1 — \"The physical security procedure shall be reviewed annually "
                         "as part of routine organisational assessments and post-incident reviews.\"",
            "system": "N/A (org-level control)",
            "severity": "Rendah (Low)",
            "condition": (
                "The Physical Security Procedure (STS-PSP-2024-001, v2.0) was last reviewed and "
                "approved on 10/11/2024. The annual review was due by 10/11/2025 but had not been "
                "completed as at the audit date (14/02/2026) — 3 months overdue. The Facilities "
                "Manager (Encik Azman, interview 30/01/2026) confirmed the review was scheduled for "
                "Q1 2026 and was in progress but not yet completed."
            ),
            "criteria": (
                "CoP 13.1.5.1 requires annual review of the physical security procedure."
            ),
            "cause": (
                "Review was deferred due to the Facilities Manager's involvement in the new data "
                "centre construction project (2025). No automated review schedule tracking was in place."
            ),
            "consequence": (
                "The physical security procedure may not reflect current threats, recent incidents, "
                "or changes to physical infrastructure (e.g., new data centre). Risk is low as the "
                "review is in progress and the existing procedure remains substantially current. "
                "No physical security incidents were reported during the period."
            ),
            "corrective": (
                "1. Complete the physical security procedure review within 30 days.\n"
                "2. Implement automated review schedule tracking for all security procedures.\n"
                "3. Ensure the updated procedure covers the new data centre physical controls."
            ),
            "legal_class": "CoP 13.1.5",
            "compoundable": "N/A — CoP non-compliance.",
        },
    ]
    add_findings_section(doc, findings)

    # --- Section 7: Directive Compliance ---
    directive_data = {
        "directive_summary": [
            ("CE Directive 01", "CSIRT Establishment", "5", "0 P / 0 SP / 5 TP / 0 TB"),
            ("CE Directive 02", "Authorised Persons Registration", "4", "1 P / 0 SP / 3 TP / 0 TB"),
            ("CE Directive 03", "Information Provision to CE", "3", "3 P / 0 SP / 0 TP / 0 TB"),
            ("CE Directive 04", "NCII System Information", "4", "4 P / 0 SP / 0 TP / 0 TB"),
            ("CE Directive 05", "Risk Assessment Process", "6", "4 P / 1 SP / 1 TP / 0 TB"),
            ("CE Directive 06", "Incident Notification", "5", "0 P / 0 SP / 5 TP / 0 TB"),
            ("CE Directive 07", "Threat Information Sharing", "3", "2 P / 0 SP / 1 TP / 0 TB"),
            ("CE Directive 08", "Audit Requirements", "5", "5 P / 0 SP / 0 TP / 0 TB"),
            ("CE Directive 09", "Compliance Monitoring", "4", "4 P / 0 SP / 0 TP / 0 TB"),
            ("CE Directive 10", "Sector-Specific Directives", "3", "2 P / 1 SP / 0 TP / 0 TB"),
        ],
        "directive_findings": [
            {
                "id": "DIR-2026-001",
                "title": "Orang Diberi Kuasa Tidak Didaftarkan dengan NACSA / "
                         "Authorised Persons Not Registered with NACSA",
                "directive": "CE Directive 02 — Authorised Persons Registration",
                "para_ref": "Directive 02, Para 6-8",
                "severity": "Kritikal (Critical)",
                "condition": (
                    "STS has not registered any authorised persons with NACSA as required by "
                    "CE Directive 02. The entity has designated internal authorised persons "
                    "(CISO, IT Security Manager, Network Operations Manager) but has not submitted "
                    "their particulars to NACSA for approval. Without registered authorised persons, "
                    "the entity cannot officially communicate with NACSA on cyber security matters "
                    "or receive classified threat intelligence."
                ),
                "criteria": (
                    "CE Directive 02 requires NCII entities to register authorised persons with "
                    "NACSA and maintain current registration at all times."
                ),
                "cause": (
                    "The CISO was unaware that authorised person registration was a separate "
                    "requirement from NCII entity designation. No internal tracking of CE Directive "
                    "compliance obligations existed prior to this audit."
                ),
                "consequence": (
                    "Without registered authorised persons, STS cannot receive classified threat "
                    "intelligence from NACSA, cannot participate in sector-level threat sharing, "
                    "and may face delays in official communications during a cyber security incident. "
                    "This compounds the CSIRT finding (CoP-2026-001) — the entity is effectively "
                    "isolated from the national cyber security ecosystem."
                ),
                "corrective": (
                    "1. Submit authorised person registration forms to NACSA immediately.\n"
                    "2. Register minimum 3 authorised persons covering 24/7 availability.\n"
                    "3. Establish process for updating registrations on personnel changes."
                ),
            },
        ],
    }
    add_directive_compliance_section(doc, directive_data)

    # --- Section 8: Domain Compliance Profile ---
    profile_data = {
        "domain_profile": [
            ("4.0", "Governance and Leadership", "1", "1", "0", "0", "0"),
            ("5.0", "Cyber Security Policy and Objectives", "11", "9", "1", "1", "0"),
            ("6.0", "Organisational Development", "19", "15", "3", "1", "0"),
            ("7.0", "Cyber Security Assurance", "24", "20", "3", "1", "0"),
            ("8.0", "Resource Allocation and Optimisation", "19", "16", "2", "1", "0"),
            ("9.0", "Risk Management", "19", "14", "3", "2", "0"),
            ("10.0", "Operational Efficiency", "9", "8", "1", "0", "0"),
            ("11.0", "Data Security", "23", "19", "3", "1", "0"),
            ("12.0", "Contractual Management", "6", "6", "0", "0", "0"),
            ("13.0", "Physical Security", "15", "13", "1", "1", "0"),
            ("14.0", "System and Network Security", "24", "18", "4", "2", "0"),
            ("15.0", "Access Control", "13", "11", "1", "1", "0"),
            ("16.0", "Technical Vulnerability", "10", "7", "2", "1", "0"),
            ("17.0", "Cyber Security Event Management", "3", "3", "0", "0", "0"),
            ("18.0", "Cyber Security Incident Management", "6", "2", "1", "3", "0"),
            ("19.0", "Business Continuity Management", "6", "4", "1", "0", "1"),
            ("20.0", "Sector-Specific Requirements", "2", "1", "0", "0", "1"),
            ("21.0", "Monitoring and Compliance", "0", "—", "—", "—", "—"),
            ("", "JUMLAH / TOTAL", "210", "167", "26", "15", "2"),
        ],
    }
    add_domain_compliance_profile(doc, profile_data)

    # --- Section 9: Finding Summary ---
    register_data = {
        "register": [
            ("CoP-2026-001", "CSIRT Not Established", "CoP 18.1.1.1", "Kritikal", "Open", "26/03/2026"),
            ("CoP-2026-002", "Risk Assessment Submission Overdue", "CoP 9.1.4.2", "Tinggi", "Open", "03/03/2026"),
            ("CoP-2026-003", "SCADA Firewall Rules Stale", "CoP 14.3.4.1", "Tinggi", "Open", "26/05/2026"),
            ("CoP-2026-004", "Vulnerability Scanning Gap", "CoP 16.2.1.1", "Sederhana", "Open", "26/08/2026"),
            ("CoP-2026-005", "DLP Gap on SCADA", "CoP 11.1.1.1", "Sederhana", "Open", "26/08/2026"),
            ("CoP-2026-006", "Physical Security Review Overdue", "CoP 13.1.5.1", "Rendah", "Open", "26/03/2026"),
            ("DIR-2026-001", "Authorised Persons Not Registered", "Directive 02", "Kritikal", "Open", "26/03/2026"),
        ],
    }
    add_finding_summary_section(doc, register_data)

    # --- Section 10: Improvement Opportunities (Tier 2) ---
    improvement_data = {
        "opportunities": [
            {
                "num": "001",
                "title": "Mekanisme Pemantauan Pematuhan Objektif / Compliance Monitoring Mechanism",
                "clause_ref": "CoP 5.2.3.2 — \"The NCII entity should have a monitoring mechanism to "
                              "validate compliance with the objectives.\" (SHOULD clause)",
                "obligation": "SHOULD (Tier 2)",
                "observation": (
                    "STS has defined cyber security objectives (CoP 5.2.1) and communicates them "
                    "to stakeholders (CoP 5.2.2). However, there is no structured monitoring mechanism "
                    "to track whether objectives are being achieved. Progress is tracked informally "
                    "through monthly CISO reports to the Board but there is no dashboard, KPI tracking, "
                    "or automated compliance monitoring tool."
                ),
                "recommendation": (
                    "Implement a cyber security objectives dashboard with measurable KPIs "
                    "(e.g., patch compliance rate, MTTD, MTTR, awareness training completion rate). "
                    "Integrate with the quarterly Board reporting cycle."
                ),
            },
            {
                "num": "002",
                "title": "Pelan Komunikasi Polisi / Policy Communication Plan",
                "clause_ref": "CoP 5.1.5.2 — \"A communication plan may be established to ensure the "
                              "Cyber Security Policy is disseminated, understood, and enforced across all "
                              "levels of the organisation.\" (MAY clause)",
                "obligation": "MAY (Tier 2)",
                "observation": (
                    "The cyber security policy is communicated through intranet and annual awareness "
                    "training. However, no formal communication plan exists to ensure systematic "
                    "dissemination to all personnel, including contractors and third-party staff "
                    "who access NCII systems."
                ),
                "recommendation": (
                    "Develop a formal communication plan for the cyber security policy covering "
                    "all personnel categories (permanent staff, contractors, third-party maintenance). "
                    "Include acknowledgement tracking."
                ),
            },
            {
                "num": "003",
                "title": "Proses Semakan Polisi Berstruktur / Structured Policy Review Process",
                "clause_ref": "CoP 5.1.4.2 — \"A structured policy review process may include... "
                              "version control and change management procedures.\" (MAY clause)",
                "obligation": "MAY (Tier 2)",
                "observation": (
                    "Policy reviews are conducted but the review process is informal — there is no "
                    "documented review checklist, no formal sign-off by review participants, and "
                    "version control relies on filename conventions rather than a document management system."
                ),
                "recommendation": (
                    "Adopt a structured policy review process with a documented checklist, "
                    "formal sign-off by all reviewers, and a document management system with "
                    "version control, audit trails, and access control."
                ),
            },
        ],
    }
    add_improvement_section(doc, improvement_data)

    # --- Section 11: Conclusion ---
    conclusion_data = {
        "posture": [
            "Of 210 Tier 1 CoP requirements assessed, 167 (79.5%) were rated Patuh, 26 (12.4%) Separa Patuh, "
            "15 (7.1%) Tidak Patuh, and 2 (1.0%) Tidak Berkenaan. Of 42 CE Directive requirements, "
            "25 (59.5%) were Patuh, 2 (4.8%) Separa Patuh, and 15 (35.7%) Tidak Patuh.",
            "Domains with strongest compliance: CoP 4.0 (Governance — 100%), CoP 12.0 (Contractual — 100%), "
            "CoP 17.0 (Event Management — 100%), CoP 10.0 (Operational Efficiency — 89%).",
            "Domains with weakest compliance: CoP 18.0 (Incident Management — 33% TP), "
            "CoP 9.0 (Risk Management — 11% TP), CoP 14.0 (System/Network Security — 8% TP).",
            "Of 11 statutory obligations, 7 were Met, 1 Pending (s23 — this audit report), "
            "1 Partial (s27), and 2 Not Met (s21, s25, s26). The s26 non-compliance is Critical.",
        ],
        "risk_summary": [
            "6 CoP findings and 1 directive finding were raised: 1 Critical (+ 1 directive Critical), "
            "2 High, 2 Medium, and 1 Low.",
            "The greatest risk to NCII is the absence of incident response capability (CSIRT). "
            "The energy SCADA system controls power distribution to ~180,000 consumers and has no "
            "structured incident detection, classification, or response capability. Combined with "
            "the stale firewall rules finding, the SCADA segment represents the highest-risk area.",
            "The combination of no CSIRT, no registered authorised persons, and untested 6-hour "
            "notification creates a scenario where a significant cyber incident could occur and "
            "NACSA would not be notified within the statutory timeframe, hampering national-level "
            "incident coordination.",
        ],
        "attention": [
            "CSIRT establishment and 6-hour notification testing (CoP-2026-001, DIR-2026-001) — "
            "Critical. Remediation within 30 days.",
            "Risk assessment report submission (CoP-2026-002) — High. Submit immediately (within 7 days).",
            "SCADA firewall rule review (CoP-2026-003) — High. Remediation within 90 days.",
            "Systemic pattern: OT/SCADA security controls are consistently weaker than IT controls "
            "across multiple domains (CoP 11.0, 14.0, 16.0). This reflects inadequate investment in "
            "OT-specific cyber security capability.",
        ],
        "prior_audit": "This is STS's first cyber security audit under Act 854. No prior findings exist.",
        "opinion_rows": [
            ("Memuaskan (Satisfactory)",
             "No Critical/High findings. TP count zero or near-zero.", ""),
            ("Perlu Penambahbaikan\n(Needs Improvement)",
             "One or more High, OR multiple Medium with systemic gaps.", "[X]"),
            ("Tidak Memuaskan (Unsatisfactory)",
             "One or more Critical, OR multiple High across core domains.", ""),
        ],
        "legal_input": [
            "The legal practitioner (Puan Siti Nurhaliza) has reviewed the statutory compliance "
            "assessment (Section 4) and confirms the factual accuracy of the s21, s25, and s26 "
            "non-compliance determinations.",
            "Penalty exposure for the entity: s26 non-compliance (failure to establish incident notification "
            "process) — fine up to RM500,000 or imprisonment up to 10 years or both. s21 non-compliance "
            "(late risk assessment submission) — fine up to RM200,000 or imprisonment up to 3 years or both "
            "(compoundable at RM100,000).",
            "s52 director liability exposure has been assessed for 4 officers (see Section 5). "
            "The CEO and CFO face the highest exposure given documented evidence that CSIRT budget "
            "was specifically requested by the CISO and not approved.",
            "Note: The full privileged penalty quantification and legal strategy (Mode B) has been "
            "delivered separately to the audit committee chair under legal privilege. It is NOT "
            "referenced in this report.",
        ],
    }
    add_conclusion_section(doc, conclusion_data)

    # --- Section 12: Management Action Plan ---
    map_data = {
        "map_rows": [
            ("CoP-2026-001", "Setuju / Agree",
             "Establish CSIRT. Budget approved by Board (emergency resolution 20/02/2026). "
             "CSIRT to be operational by 26/03/2026. NC4 portal activation in progress.",
             "CISO", "26/03/2026", "In Progress"),
            ("CoP-2026-002", "Setuju / Agree",
             "CEO signed risk assessment report on 18/02/2026. Submission to "
             "auditpematuhan@nacsa.gov.my completed on 19/02/2026.",
             "CISO", "19/02/2026", "Closed"),
            ("CoP-2026-003", "Setuju / Agree",
             "Emergency firewall rule review for SCADA initiated. External OT security "
             "vendor (Dragos) engaged for comprehensive review. Quarterly review schedule "
             "to be reinstated.",
             "CTO", "26/05/2026", "In Progress"),
            ("CoP-2026-004", "Setuju / Agree",
             "Passive scanning solution for OT assets being evaluated. RFQ issued to "
             "3 vendors. Interim: firmware version tracking for all excluded assets "
             "implemented manually.",
             "CISO", "26/08/2026", "In Progress"),
            ("CoP-2026-005", "Sebahagian / Partial",
             "Entity agrees DLP coverage should be extended but notes performance constraints "
             "on SCADA HMI workstations. Will implement network-based DLP on SCADA/IT boundary "
             "as compensating control. Full endpoint DLP for SCADA not feasible.",
             "CTO", "26/08/2026", "In Progress"),
            ("CoP-2026-006", "Setuju / Agree",
             "Physical security procedure review completed on 10/02/2026 (during audit). "
             "Updated procedure covers new data centre. Automated review tracking implemented "
             "in ServiceNow.",
             "Facilities Mgr", "10/02/2026", "Closed"),
            ("DIR-2026-001", "Setuju / Agree",
             "Authorised person registration forms prepared for CISO, IT Security Manager, "
             "and NOC Manager. Submission to NACSA by 28/02/2026.",
             "CISO", "28/02/2026", "In Progress"),
        ],
    }
    add_map_section(doc, map_data)

    # --- Appendices ---
    appendix_data = {
        "systems": [
            ("NCII-SYS-001", "Sistem Kawalan SCADA / SCADA Control System",
             "ABB Ability Symphony Plus. Controls power distribution for Northern Region grid. "
             "142 endpoints (PLCs, RTUs, HMIs).",
             "Kritikal", "Pusat Kawalan Utara / Northern Control Centre, Penang", "VP Operations"),
            ("NCII-SYS-002", "Sistem Pengurusan Tenaga / Energy Management System (EMS)",
             "GE EMS/SCADA. Real-time monitoring of generation, transmission, and distribution. "
             "Interfaces with SCADA via ICCP.",
             "Kritikal", "Pusat Kawalan Utara / Northern Control Centre, Penang", "VP Operations"),
            ("NCII-SYS-003", "Platform Grid Pintar / Smart Grid Platform",
             "Custom-built platform for smart meter management, demand response, and "
             "distributed energy resource integration. 45,000 smart meters.",
             "Standard", "Data Centre Kuala Lumpur / KL Data Centre", "CTO"),
            ("NCII-SYS-004", "Rangkaian Korporat / Corporate Network",
             "Enterprise IT network supporting business operations, ERP, email, and "
             "remote access. Interconnected with SCADA via controlled DMZ.",
             "Standard", "Ibu Pejabat / Head Office, Kuala Lumpur", "CTO"),
        ],
        "penalties": [
            ("s26", "Failure to notify cyber security incident within 6 hours",
             "RM500,000", "10 years", "No",
             "CoP-2026-001"),
            ("s21", "Failure to submit risk assessment report within prescribed period",
             "RM200,000", "3 years", "Yes (RM100,000)",
             "CoP-2026-002"),
            ("s25", "Failure to share cyber threat information when required",
             "RM200,000", "3 years", "Yes (RM100,000)",
             "DIR-2026-001"),
            ("s52", "Director liability — deemed to have committed offence unless due diligence proved",
             "Same as principal offence", "Same as principal offence", "Same as principal offence",
             "CoP-2026-001, CoP-2026-002"),
        ],
    }
    add_appendices(doc, appendix_data)

    # --- Sign-off ---
    signoff_data = {
        "signoffs": [
            ("Disediakan oleh / Prepared by:", "Ketua Juruaudit / Lead Auditor",
             "Ahmad Faizal bin Mohd Noor", "24/02/2026"),
            ("Disemak oleh / Reviewed by:", "Rakan Kongsi Jaminan / Engagement Partner",
             "Dato' Ir. Razlan Hashim", "24/02/2026"),
            ("Semakan undang-undang / Legal review (Mode A):",
             "Pengamal Undang-Undang / Legal Practitioner",
             "Puan Siti Nurhaliza binti Ahmad", "23/02/2026"),
            ("Pengesahan entiti / Entity confirmation:",
             "Pegawai Diberi Kuasa / Authorised Officer — CEO",
             "Encik Kamal bin Yusoff", "24/02/2026"),
            ("Dihantar kepada NACSA / Submitted to NACSA:",
             "Via auditpematuhan@nacsa.gov.my",
             "Salinan 1 / Copy 1", "26/02/2026"),
        ],
    }
    add_signoff(doc, signoff_data)

    doc.save(output_path)
    print(f"Sample report saved: {output_path}")


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    # --- Document 1: Report Template ---
    template_path = os.path.join(SCRIPT_DIR, "NACSA-Audit-Report-Template.docx")
    build_report_template(template_path)

    # --- Document 2: Sample Report ---
    sample_path = os.path.join(SCRIPT_DIR, "NACSA-Sample-Audit-Report.docx")
    build_sample_report(sample_path)

    # --- Copy to Tech-Audit/NACSA ---
    tech_audit = os.path.expanduser("~/claude/Tech-Audit/NACSA")
    if os.path.isdir(tech_audit):
        import shutil
        for src_name in ["NACSA-Audit-Report-Template.docx", "NACSA-Sample-Audit-Report.docx"]:
            src = os.path.join(SCRIPT_DIR, src_name)
            dst = os.path.join(tech_audit, src_name)
            shutil.copy2(src, dst)
            print(f"Copied to: {dst}")
    else:
        print(f"Tech-Audit directory not found: {tech_audit}")
